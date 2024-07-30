import logging
import time
from .ollama.request_handler import llm_request_handler
from django.core.files.uploadedfile import UploadedFile
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class SummarizePfHandler:
    def __init__(self, files: list[UploadedFile]):
        self.files = files
        self.request_manager = llm_request_handler()

    def extract_text_from_docx(self, file: UploadedFile):
        start_time = time.time()
        document = Document(file)
        full_text = []
        seen_elements = set()  # To keep track of seen elements and avoid duplicates

        for para in document.paragraphs:
            element_text = para.text.strip()
            if element_text and element_text not in seen_elements:
                full_text.append(element_text)
                seen_elements.add(element_text)
        
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text not in seen_elements:
                        row_text.append(cell_text)
                        seen_elements.add(cell_text)
                if row_text:
                    full_text.append('\t'.join(row_text))

        extracted_text = '\n\n'.join(full_text)

        # Clean repeated sections
        cleaned_text = self.clean_repeated_sections(extracted_text)

        extraction_time = time.time() - start_time
        logging.info(f"Time to extract text from {file.name}: {extraction_time:.2f} seconds")

        return cleaned_text

    def clean_repeated_sections(self, text):
        # Split text into lines
        lines = text.split('\n')
        
        # Remove consecutive duplicate lines
        cleaned_lines = []
        previous_line = ""
        for line in lines:
            if line.strip() != previous_line.strip():
                cleaned_lines.append(line)
                previous_line = line

        # Further remove duplicate sections by splitting by double newline and checking uniqueness
        unique_sections = []
        seen_sections = set()
        for section in '\n\n'.join(cleaned_lines).split('\n\n'):
            if section not in seen_sections:
                unique_sections.append(section)
                seen_sections.add(section)

        return '\n\n'.join(unique_sections)


    def summarize_pf_from_files(self, language):
        all_texts = []

        # Extracts text from the files and concatenates them
        for file in self.files:
            text = self.extract_text_from_docx(file)
            all_texts.append(text)

        combined_text = '\n\n'.join(all_texts)

        prompt = (
            f"""You are a text summarization and medical expert. Given the following medical record, generate a medical summary in {language} in the following format (replace the values between "" by what you found in the medical record):

            Record:
            {combined_text}

            Format:
            Données personnelles:
            - Prénom: "Prénom"
            - Nom : "Nom"
            - Sexe: "Sexe"
            - N°AVS: "N°AVS"
            - Date de naissance: "Date de naissance"
            - Adresse: "Adresse"

            Cas:
            ""Complete sentence describing the case""

            Réseau social:
            - Référent: "Référent", tel: "Tel référent"
            - Médecin traitant: "Nom du médecin traitant", tel: "Tel médecin"

            Informations générales:
            "Complete sentence"

            Allergies:
            - "List of allergies"

            Historique médical:
            "Complete senetence describing the "Diagnostics""

            Médication:
            - Voie orale:
                - "List of medicaments"
            - Injection:
                - "Médicaments injection"
            """
            )

        start_time = time.time()
        response = self.request_manager.generate_response(prompt=prompt)
        summary = ''.join(response)
        generation_time = time.time() - start_time
        logging.info(f"Time to generate summary: {generation_time:.2f} seconds")

        return summary
